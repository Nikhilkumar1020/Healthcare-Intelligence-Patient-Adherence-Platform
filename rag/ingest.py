"""
rag/ingest.py
Ingests knowledge base documents into ChromaDB vector store.

Pipeline:
  Text files → text extraction → chunking → embedding → ChromaDB

Usage:
    python rag/ingest.py
"""
import sys
import os
import logging
from pathlib import Path
from typing import List, Dict

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
from dotenv import load_dotenv
load_dotenv()

logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(message)s")
logger = logging.getLogger(__name__)

DOCS_DIR     = Path(__file__).parent / "documents"
CHROMA_DIR   = Path(__file__).parent / "chroma_db"
CHUNK_SIZE   = int(os.getenv("RAG_CHUNK_SIZE", 500))
CHUNK_OVERLAP = int(os.getenv("RAG_CHUNK_OVERLAP", 50))

# Using HuggingFace embeddings for Groq compatibility


def load_documents() -> List[Dict[str, str]]:
    """Load all .txt and .pdf files from the documents directory."""
    docs = []
    for ext in ["*.txt", "*.pdf", "*.docx"]:
        for fpath in DOCS_DIR.glob(ext):
            try:
                if fpath.suffix == ".txt":
                    content = fpath.read_text(encoding="utf-8")
                elif fpath.suffix == ".pdf":
                    import pypdf
                    reader = pypdf.PdfReader(str(fpath))
                    content = "\n".join(page.extract_text() or "" for page in reader.pages)
                elif fpath.suffix == ".docx":
                    import docx
                    doc = docx.Document(str(fpath))
                    content = "\n".join(p.text for p in doc.paragraphs)
                else:
                    continue

                docs.append({"filename": fpath.name, "content": content})
                logger.info(f"[ingest] Loaded: {fpath.name} ({len(content)} chars)")
            except Exception as e:
                logger.error(f"[ingest] Failed to load {fpath.name}: {e}")
    return docs


def chunk_text(text: str, chunk_size: int, overlap: int) -> List[str]:
    """Simple sliding-window chunker."""
    words = text.split()
    chunks = []
    start = 0
    while start < len(words):
        end = start + chunk_size
        chunk = " ".join(words[start:end])
        chunks.append(chunk)
        start += chunk_size - overlap
    return chunks


def build_vector_store(docs: List[Dict[str, str]]) -> None:
    """Build ChromaDB vector store from documents."""
    CHROMA_DIR.mkdir(parents=True, exist_ok=True)

    # Use local HuggingFace embeddings when using Groq
    logger.info("[ingest] Using sentence-transformers embeddings")
    try:
        from langchain_community.embeddings import HuggingFaceEmbeddings
        embeddings = HuggingFaceEmbeddings(
            model_name="sentence-transformers/all-MiniLM-L6-v2"
        )
    except Exception as e:
        logger.error(f"[ingest] No embedding model available: {e}")
        return

    from langchain_chroma import Chroma
    from langchain_core.documents import Document

    all_chunks = []
    for doc in docs:
        chunks = chunk_text(doc["content"], CHUNK_SIZE, CHUNK_OVERLAP)
        for i, chunk in enumerate(chunks):
            all_chunks.append(Document(
                page_content=chunk,
                metadata={"source": doc["filename"], "chunk": i}
            ))

    logger.info(f"[ingest] Total chunks: {len(all_chunks)}")

    # Create/overwrite ChromaDB
    vectorstore = Chroma.from_documents(
        documents=all_chunks,
        embedding=embeddings,
        persist_directory=str(CHROMA_DIR),
        collection_name="healthcare_kb",
    )
    logger.info(f"[ingest] Vector store built at {CHROMA_DIR}")


def ingest() -> None:
    docs = load_documents()
    if not docs:
        logger.error("[ingest] No documents found. Check rag/documents/")
        return
    build_vector_store(docs)
    logger.info(f"[ingest] Ingestion complete. {len(docs)} documents processed.")


if __name__ == "__main__":
    ingest()
